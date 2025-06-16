from text_app.models import TblText, TblSentence, TblToken, TblMarkup, TblTokenMarkup
from user_app.models import TblUser, TblStudent
from generator_app.models import TblTask, TblAdditionalVariant, TblTest, TblAssignedTest, TblUserAnswer
from django.core.exceptions import ObjectDoesNotExist
from .task_controller import save_task, get_task, delete_task, check_answer
from datetime import datetime
from django.forms.models import model_to_dict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Возвращает тест по его ID.
def get_test(test_id):
    try:
        test = model_to_dict(TblTest.objects.get(test_id=test_id))
        test["tasks"] = [
            # Еее, дублирование запросов ради недублирования кода, еее.
            # Работает приемлемо и ладно.
            get_task(task["task_id"])
            for task in TblTask.objects.filter(test_id=test_id).all().values()
        ]
    except ObjectDoesNotExist:
        return None

    return test	

# Возвращает тесты, доступные студенту, по его ID.
def get_tests_id_by_stud(stud_id):
    assigns = [
        assign["test_id_id"]
        for assign in TblAssignedTest.objects.filter(user_id=stud_id).all().values()
    ]
    return [
        get_test(test["test_id"])
        for test in list((TblTest.objects.filter(creator=stud_id) | TblTest.objects.filter(test_id__in=assigns)).all().values())
    ]

# Возвращает студентов по тесту.
def get_studs_by_test(test_id):
    assigns = [
        assign["user_id_id"]
        for assign in TblAssignedTest.objects.filter(test_id=test_id).all().values()
    ]
    studs = [
        dict(model_to_dict(stud), fullname=str(stud))
        for stud in TblStudent.objects.filter(user_id__in=assigns).all()
    ]

    return studs


# Возвращает тесты, созданные преподавателем, по его ID.
def get_tests_id_by_prof(prof_id):
    # Это не халтура, это "грамотное переиспользование кода" 🤡
    return [
        get_test(test["test_id"])
        for test in list(TblTest.objects.filter(creator=prof_id).all().values())
    ]


# Возвращает информацию о состоянии решения упражнения теста студентом.
def get_task_status_by_test_and_stud(task_id, test_id, stud_id):
    try:
        answer = TblUserAnswer.objects.get(user_id=stud_id, test_id=test_id, task_id=task_id)
    except ObjectDoesNotExist:
        return None

    return model_to_dict(answer)	


# Возвращает информацию о состоянии решения упражнений теста студентом.
def get_test_status_by_stud(test_id, stud_id):
    return [
        answer
        for answer in TblUserAnswer.objects.filter(user_id=stud_id, test_id=test_id).all().values()
    ]


# Принимает на вход ответ студента. Сохраняет информацию о состоянии решения упражнения теста студентом.
def save_task_status_by_test_and_stud(task_id, test_id, stud_id, stud_answer):
    try:
        answer = TblUserAnswer.objects.get(user_id=stud_id, test_id=test_id, task_id=task_id)
    except ObjectDoesNotExist:
        answer = TblUserAnswer.objects.create(
            user_id     = TblUser.objects.get(id_user=stud_id),
            test_id     = TblTest.objects.get(test_id=test_id),
            task_id     = TblTask.objects.get(task_id=task_id),
            user_input  = stud_answer,
        )
        answer.save()
    print(answer)
    
    return model_to_dict(answer)

# Высчитывает долю правильно решённых заданий теста.
def get_test_solved_ratio(test_id, stud_id):
    test = get_test(test_id)
    answers = get_test_status_by_stud(test_id, stud_id)
    tasks = test["tasks"]
    
    for task in tasks:
        for ans in answers:
            if ans["task_id_id"] is task["task_id"]:
                task["ans"] = ans["user_input"]

    total_am = len(tasks)
    right_am = 0
    for task in tasks:
        if check_answer(task["task_id"], task.get("ans", "")):
            right_am += 1
    
    ratio = right_am / total_am
    return ratio

# Помечает тест как начатый.
def start_test(test_id, stud_id):
    assign = TblAssignedTest.objects.get(user_id=stud_id, test_id=test_id)

    if assign.start_date is None:
        assign.start_date = datetime.now()
        assign.save()

    return model_to_dict(assign) 

# Помечает тест как завершённый, вычисляет и сохраняет оценку.
def complete_test(test_id, stud_id):
    assign = TblAssignedTest.objects.get(user_id=stud_id, test_id=test_id)

    if assign.finish_date is None:
        test = get_test(test_id)
        
        ratio = get_test_solved_ratio(test_id, stud_id)
        mark = 2

        if ratio >= test["score_for_5"]:
            mark = 5
        elif ratio >= test["score_for_4"]:
            mark = 4
        elif ratio >= test["score_for_3"]:
            mark = 3

        assign.score = mark
        assign.finish_date = datetime.now()
        assign.save()

    return model_to_dict(assign) 
    

# Создаёт тест на основе переданных данных (список упражнений и параметры теста). Возвращает созданный тест с присвоенным ID.
def create_test(data):
    test = TblTest.objects.create(
        creator     = TblUser.objects.get(id_user=data["creator"]),
        name        = data["name"],
        create_date = datetime.now(),
        score_for_3 = data["score_for_3"],
        score_for_4 = data["score_for_4"],
        score_for_5 = data["score_for_5"],
    )
    print(test)
    test.save()

    tasks = []
    for task in data["tasks"]:
        task["test_id"] = test
        tasks.append(save_task(task))
    
    test = model_to_dict(test)
    test["tasks"] = tasks

    return test


# Добавляет запись о том, что данный тест отправлен студенту.
def send_test_to_stud(test_id, stud_id):
    try:
        assign = TblAssignedTest.objects.get(user_id=stud_id, test_id=test_id)
    except ObjectDoesNotExist:
        test = TblTest.objects.get(test_id=test_id)
        user = TblUser.objects.get(id_user=stud_id)
        assign = TblAssignedTest.objects.create(
            user_id = user,
            test_id = test,
        )

    return model_to_dict(assign)


# Проверяет запись о том, что данный тест назначен студенту.
def is_test_assigned(test_id, stud_id):
    try:
        assign = TblAssignedTest.objects.get(user_id=stud_id, test_id=test_id)
    except ObjectDoesNotExist:
        return False

    return True


# Отменяет запись о том, что данный тест отправлен студенту.
def unsend_test_to_stud(test_id, stud_id):
    TblUserAnswer.objects.filter(test_id=test_id, user_id=stud_id).delete()
    TblAssignedTest.objects.filter(test_id=test_id, user_id=stud_id).delete()


# Преобразует тест по ID в DOCX формат, доступный для скачивания.
def generate_docx(test_id):
    test = get_test(test_id)
    document = Document()
    document.add_heading(test['name'], 0)
    task_num = 1
    for task in test['tasks']:
        if task['input_type'] == 1:
            str1 = str(task_num)
            document.add_paragraph(str1 + ". " + task['text_before'] + " _________ " + task['text_after'])
            task_num+=1
            p = document.add_paragraph('Варианты ответа: ')
            int = 1
            for variants in task['variants']:
                str2 = str(int)
                p.add_run(str2 + ") " + variants['variant_text'] + ' ')
                int += 1
        elif task['input_type'] == 0:
            str1 = str(task_num)
            document.add_paragraph(str1 + ". " + task['text_before'] + " _________ " + task['text_after'] + ' (' + task['inf'] + ') ')
            task_num += 1
    return document


# Удаляет тест по его ID.
def delete_test(test_id):
    tasks = [
        task["task_id"]
        for task in TblTask.objects.filter(test_id=test_id).all().values()
    ]

    TblUserAnswer.objects.filter(test_id=test_id).delete()
    
    for task in tasks:
        delete_task(task)

    TblAssignedTest.objects.filter(test_id=test_id).delete()
    
    TblTest.objects.filter(test_id=test_id).delete()


def generate_report(stud_ids, test_ids):

    solves = (TblAssignedTest.objects.filter(test_id__in=test_ids) & TblAssignedTest.objects.filter(user_id__in=stud_ids) & TblAssignedTest.objects.filter(score__isnull=False)).all().values()
    users = TblUser.objects.filter(id_user__in=stud_ids).all().values()
    users_ids = {}
    for user in users:
        users_ids[user["id_user"]] = user["name"] + ' ' + user["last_name"] + ' ' + (user["patronymic"] or "")

    print(users_ids)
    tests_solves = {}
    for test_id in test_ids:
        tests_solves[test_id] = []
    
    for solve in solves:
        solve["user_name"] = users_ids[solve["user_id_id"]]
        tests_solves[str(solve["test_id_id"])].append(solve)

    print(solves)
    
    document = Document()

    for test_id, test_solves in tests_solves.items():
        if len(test_solves) < 1:
            continue
        
        test = get_test(test_id)
        
        document.add_heading(test['name'])
        table = document.add_table(1, 4)
        table.style = 'Light Shading Accent 1'
        head_cells = table.rows[0].cells

        for i, item in enumerate(['ФИО студента', 'Дата решения', 'Процент верно решенных упражнений', 'Оценка']):
            p = head_cells[i].paragraphs[0]
            p.add_run(item).bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for solve in test_solves:
            cells = table.add_row().cells
            cells[0].text = str(solve["user_name"])
            cells[1].text = str(solve["finish_date"].strftime("%d.%m.%y %H:%M:%S"))
            cells[2].text = str(get_test_solved_ratio(solve["test_id_id"], solve["user_id_id"]))
            cells[3].text = str(solve["score"])

            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        document.add_paragraph()
        

    return document